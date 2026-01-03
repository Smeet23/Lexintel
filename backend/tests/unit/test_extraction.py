import pytest
from app.services.extraction import chunk_text


def test_chunk_text_basic():
    """Test basic text chunking"""
    text = "a" * 10000  # 10,000 characters
    chunks = chunk_text(text, chunk_size=4000, overlap=400)

    assert len(chunks) == 3  # 4000 + (4000-400) + remainder
    assert chunks[0] == "a" * 4000
    assert len(chunks[1]) == 4000
    assert chunks[1][:400] == "a" * 400  # Overlap from previous


def test_chunk_text_small():
    """Test chunking with text smaller than chunk size"""
    text = "hello world"
    chunks = chunk_text(text, chunk_size=4000, overlap=400)

    assert len(chunks) == 1
    assert chunks[0] == "hello world"


def test_chunk_text_empty():
    """Test chunking empty text"""
    text = ""
    chunks = chunk_text(text, chunk_size=4000, overlap=400)

    assert len(chunks) == 0


def test_chunk_text_exact_chunk_size():
    """Test chunking text exactly divisible by chunk size"""
    text = "a" * 8000
    chunks = chunk_text(text, chunk_size=4000, overlap=400)

    assert len(chunks) == 3
    assert chunks[0] == "a" * 4000
    assert chunks[1][:400] == "a" * 400  # Overlap
    assert len(chunks[2]) == 800  # Remaining


# PDF Extraction Tests
def test_extract_pdf_text(tmp_path):
    """Test extracting text from PDF"""
    from app.services.extraction import extract_pdf
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    # Create test PDF
    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "Test PDF Content")
    c.drawString(100, 730, "Second line of text")
    c.save()

    text = extract_pdf(str(pdf_path))

    assert isinstance(text, str)
    assert len(text) > 0
    assert "Test PDF Content" in text or "test" in text.lower()


def test_extract_pdf_not_found():
    """Test extraction with non-existent file"""
    from app.services.extraction import extract_pdf

    with pytest.raises(FileNotFoundError):
        extract_pdf("/nonexistent/file.pdf")


# DOCX Extraction Tests
def test_extract_docx_text(tmp_path):
    """Test extracting text from DOCX"""
    from app.services.extraction import extract_docx
    from docx import Document

    # Create test DOCX
    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("Test DOCX Content")
    doc.add_paragraph("Second paragraph")
    doc.save(str(docx_path))

    text = extract_docx(str(docx_path))

    assert isinstance(text, str)
    assert len(text) > 0
    assert "Test DOCX Content" in text
    assert "Second paragraph" in text


def test_extract_docx_not_found():
    """Test DOCX extraction with non-existent file"""
    from app.services.extraction import extract_docx

    with pytest.raises(FileNotFoundError):
        extract_docx("/nonexistent/file.docx")


# TXT Extraction Tests
def test_extract_txt_text(tmp_path):
    """Test extracting text from TXT"""
    from app.services.extraction import extract_txt

    txt_path = tmp_path / "test.txt"
    txt_path.write_text("Test TXT Content\nSecond line\nThird line")

    text = extract_txt(str(txt_path))

    assert isinstance(text, str)
    assert "Test TXT Content" in text
    assert "Second line" in text


def test_extract_txt_empty(tmp_path):
    """Test extracting from empty TXT"""
    from app.services.extraction import extract_txt

    txt_path = tmp_path / "empty.txt"
    txt_path.write_text("")

    text = extract_txt(str(txt_path))

    assert text == ""


def test_extract_txt_not_found():
    """Test TXT extraction with non-existent file"""
    from app.services.extraction import extract_txt

    with pytest.raises(FileNotFoundError):
        extract_txt("/nonexistent/file.txt")


# Universal Extract Function Tests
def test_extract_file_by_extension_pdf(tmp_path):
    """Test extract_file dispatches to correct extractor for PDF"""
    from app.services.extraction import extract_file
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "PDF Test")
    c.save()

    text = extract_file(str(pdf_path))
    assert isinstance(text, str)
    assert len(text) > 0


def test_extract_file_by_extension_docx(tmp_path):
    """Test extract_file dispatches to correct extractor for DOCX"""
    from app.services.extraction import extract_file
    from docx import Document

    docx_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_paragraph("DOCX Test")
    doc.save(str(docx_path))

    text = extract_file(str(docx_path))
    assert isinstance(text, str)
    assert "DOCX Test" in text


def test_extract_file_by_extension_txt(tmp_path):
    """Test extract_file dispatches to correct extractor for TXT"""
    from app.services.extraction import extract_file

    txt_path = tmp_path / "test.txt"
    txt_path.write_text("TXT Test Content")

    text = extract_file(str(txt_path))
    assert "TXT Test Content" in text


def test_extract_file_unsupported(tmp_path):
    """Test extract_file with unsupported file type"""
    from app.services.extraction import extract_file

    # Create a dummy .doc file (old format)
    tmp_file = tmp_path / "test.doc"
    tmp_file.write_text("dummy")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_file(str(tmp_file))


# Document Chunk Service Tests
@pytest.mark.asyncio
async def test_create_text_chunks(db_session):
    """Test creating text chunks in database"""
    from app.services.extraction import create_text_chunks

    text = "a" * 10000  # 10,000 characters
    document_id = "doc_123"

    chunk_ids = await create_text_chunks(
        document_id=document_id,
        text=text,
        session=db_session,
        chunk_size=4000,
        overlap=400,
    )

    assert len(chunk_ids) == 3  # Should create 3 chunks
    assert all(isinstance(cid, str) for cid in chunk_ids)


@pytest.mark.asyncio
async def test_create_text_chunks_empty(db_session):
    """Test creating chunks from empty text"""
    from app.services.extraction import create_text_chunks

    document_id = "doc_456"

    chunk_ids = await create_text_chunks(
        document_id=document_id,
        text="",
        session=db_session,
        chunk_size=4000,
        overlap=400,
    )

    assert len(chunk_ids) == 0


@pytest.mark.asyncio
async def test_create_text_chunks_small_text(db_session):
    """Test creating chunks from text smaller than chunk size"""
    from app.services.extraction import create_text_chunks

    text = "This is a small document"
    document_id = "doc_789"

    chunk_ids = await create_text_chunks(
        document_id=document_id,
        text=text,
        session=db_session,
        chunk_size=4000,
        overlap=400,
    )

    assert len(chunk_ids) == 1
    assert len(chunk_ids) > 0
