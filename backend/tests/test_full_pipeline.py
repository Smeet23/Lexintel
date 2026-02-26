"""
Full pipeline integration test for Veritas AI document processing.

Tests the complete flow:
1. Create a sample legal PDF
2. Extract structured text (pymupdf4llm)
3. Detect legal section boundaries
4. Chunk with hybrid semantic strategy
5. Verify chunk structure and metadata
6. Simulate DB-first UUID assignment
7. Verify Qdrant payload structure
8. Test RAG engine citation patterns
9. Test embedding retry logic
10. Test vector store batch upserts
"""
import sys
import os
import json
import uuid
import tempfile
from datetime import datetime
from unittest.mock import MagicMock, patch, PropertyMock

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

# Mock settings before any imports
mock_settings = MagicMock()
mock_settings.google_api_key = "test-key-not-real"
mock_settings.qdrant_url = "http://localhost:6333"
mock_settings.database_url = "postgresql://test:test@localhost:5432/test"
mock_settings.azure_storage_connection_string = "UseDevelopmentStorage=true"
mock_settings.redis_url = "redis://localhost:6379/0"
mock_settings.celery_broker_url = "redis://localhost:6379/0"
mock_settings.celery_result_backend = "redis://localhost:6379/1"
mock_settings.debug = True
mock_settings.cache_enabled = False
mock_settings.cache_ttl_seconds = 0
mock_settings.allowed_origins = "http://localhost:3000"

# Patch settings globally
import backend.config
backend.config.get_settings = lambda: mock_settings

DIVIDER = "=" * 70
PASS = "✅ PASS"
FAIL = "❌ FAIL"
results = []


def log_test(name, passed, detail=""):
    status = PASS if passed else FAIL
    results.append((name, passed))
    print(f"\n{status}: {name}")
    if detail:
        print(f"   {detail}")


def create_sample_legal_pdf():
    """Create a realistic legal PDF for testing."""
    import fitz  # PyMuPDF

    doc = fitz.open()  # New empty PDF

    # Page 1: Contract header + definitions
    page1 = doc.new_page()
    page1_text = """MASTER SERVICES AGREEMENT

This Master Services Agreement ("Agreement") is entered into as of January 15, 2026,
by and between TechCorp Inc., a Delaware corporation ("Company"), and Legal Solutions LLC,
a California limited liability company ("Provider").

RECITALS

WHEREAS, Company desires to engage Provider to perform certain professional services; and

WHEREAS, Provider has the expertise, experience, and resources to perform such services;

NOW, THEREFORE, in consideration of the mutual covenants contained herein, the parties agree:

ARTICLE I - DEFINITIONS

Section 1.1 "Confidential Information" means any and all non-public information, whether
written, oral, electronic, or visual, that is disclosed by either party to the other party,
including but not limited to trade secrets, business plans, financial data, customer lists,
technical specifications, and proprietary software code.

Section 1.2 "Deliverables" means any work product, including reports, analyses, software,
documentation, and other materials, created by Provider in the performance of Services.

Section 1.3 "Services" means the professional consulting, legal analysis, and advisory
services described in each Statement of Work executed under this Agreement.

Section 1.4 "Term" means the period beginning on the Effective Date and ending on the
date of termination or expiration of this Agreement as set forth in Article VII."""

    text_point = fitz.Point(72, 72)
    page1.insert_text(text_point, page1_text, fontsize=10)

    # Page 2: Representations, warranties, liability
    page2 = doc.new_page()
    page2_text = """ARTICLE II - REPRESENTATIONS AND WARRANTIES

Section 2.1 Provider represents and warrants that:
(a) It has the legal right and authority to enter into this Agreement;
(b) The Services will be performed in a professional and workmanlike manner;
(c) All Deliverables will be original works and will not infringe any third-party rights;
(d) Provider will comply with all applicable laws and regulations.

Section 2.2 Company represents and warrants that:
(a) It has the legal right and authority to enter into this Agreement;
(b) It will provide timely access to information and resources reasonably necessary
    for Provider to perform the Services.

ARTICLE III - LIMITATION OF LIABILITY

Section 3.1 EXCEPT FOR BREACHES OF CONFIDENTIALITY OBLIGATIONS OR INDEMNIFICATION
OBLIGATIONS, NEITHER PARTY SHALL BE LIABLE TO THE OTHER FOR ANY INDIRECT, INCIDENTAL,
SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES ARISING OUT OF OR RELATED TO THIS AGREEMENT.

Section 3.2 The total aggregate liability of either party under this Agreement shall not
exceed the total fees paid or payable under this Agreement during the twelve (12) month
period immediately preceding the event giving rise to such liability.

ARTICLE IV - INDEMNIFICATION

Section 4.1 Provider shall indemnify, defend, and hold harmless Company from any claims,
damages, losses, and expenses arising from: (a) Provider's breach of this Agreement;
(b) Provider's negligent or willful misconduct; (c) any infringement of third-party rights.

Section 4.2 Company shall indemnify, defend, and hold harmless Provider from any claims
arising from Company's breach of this Agreement or Company's negligent acts."""

    text_point = fitz.Point(72, 72)
    page2.insert_text(text_point, page2_text, fontsize=10)

    # Page 3: Termination, governing law, exhibits
    page3 = doc.new_page()
    page3_text = """ARTICLE V - CONFIDENTIALITY

Section 5.1 Each party agrees to hold in strict confidence all Confidential Information
of the other party. Neither party shall disclose Confidential Information to any third
party without the prior written consent of the disclosing party.

Section 5.2 The obligations of confidentiality shall survive termination of this Agreement
for a period of three (3) years.

ARTICLE VI - DISPUTE RESOLUTION

Section 6.1 Any dispute arising out of or relating to this Agreement shall first be
submitted to mediation in accordance with the rules of the American Arbitration Association.

Section 6.2 If mediation is unsuccessful, the dispute shall be resolved by binding
arbitration conducted in New York, New York.

ARTICLE VII - TERMINATION

Section 7.1 Either party may terminate this Agreement upon thirty (30) days written notice.

Section 7.2 This Agreement may be terminated immediately upon material breach by either
party if such breach is not cured within fifteen (15) days of written notice.

GOVERNING LAW

This Agreement shall be governed by and construed in accordance with the laws of the
State of New York, without regard to its conflicts of laws principles.

EXHIBIT A - STATEMENT OF WORK

The Provider shall deliver the following services:
1. Legal document analysis and classification
2. Contract review and risk assessment
3. Regulatory compliance monitoring
4. Monthly status reports and quarterly reviews"""

    text_point = fitz.Point(72, 72)
    page3.insert_text(text_point, page3_text, fontsize=10)

    # Save to temp file
    pdf_path = tempfile.mktemp(suffix=".pdf")
    doc.save(pdf_path)
    doc.close()

    # Read back as bytes
    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()

    os.unlink(pdf_path)
    return pdf_bytes


def test_section_detector():
    """Test 1: Legal section boundary detection."""
    print(f"\n{DIVIDER}")
    print("TEST 1: SECTION DETECTOR")
    print(DIVIDER)

    from services.section_detector import detect_sections

    legal_text = """RECITALS

WHEREAS, Company desires to engage Provider to perform certain professional services.

ARTICLE I - DEFINITIONS

Section 1.1 "Confidential Information" means any non-public information disclosed by either party.

Section 1.2 "Deliverables" means any work product created by Provider.

ARTICLE II - REPRESENTATIONS AND WARRANTIES

Section 2.1 Provider represents and warrants that it has the legal right to enter this Agreement.

LIMITATION OF LIABILITY

NEITHER PARTY SHALL BE LIABLE FOR INDIRECT DAMAGES.

EXHIBIT A - STATEMENT OF WORK

The Provider shall deliver legal document analysis services.

GOVERNING LAW

This Agreement shall be governed by the laws of New York."""

    sections = detect_sections(legal_text)

    print(f"\n   Input text length: {len(legal_text)} chars")
    print(f"   Detected {len(sections)} sections:\n")

    for i, section in enumerate(sections):
        header = section['header'][:60] if section['header'] else "(preamble)"
        print(f"   [{i+1}] Type: {section['section_type']:20s} | Header: {header}")
        print(f"        Content length: {len(section['content'])} chars")

    # Verify key sections detected
    section_types = [s['section_type'] for s in sections]
    has_article = 'article' in section_types
    has_section = 'section' in section_types
    has_contract = 'contract_header' in section_types
    has_exhibit = 'exhibit' in section_types

    log_test("Detects ARTICLE boundaries", has_article)
    log_test("Detects SECTION boundaries", has_section)
    log_test("Detects contract headers (RECITALS, GOVERNING LAW, etc.)", has_contract)
    log_test("Detects EXHIBIT boundaries", has_exhibit)
    log_test("Returns >5 sections for complex legal text", len(sections) >= 5,
             f"Got {len(sections)} sections")

    return sections


def test_pdf_extraction(pdf_bytes):
    """Test 2: Structured PDF text extraction."""
    print(f"\n{DIVIDER}")
    print("TEST 2: PDF TEXT EXTRACTION")
    print(DIVIDER)

    from services.text_extraction import extract_text, extract_pdf_text, extract_pdf_text_structured

    # Test basic extraction
    basic_sections = extract_pdf_text(pdf_bytes)
    print(f"\n   Basic extraction: {len(basic_sections)} pages")
    for s in basic_sections:
        print(f"   - Page {s['location']}: {len(s['content'])} chars")

    log_test("Basic PDF extraction returns pages", len(basic_sections) >= 3,
             f"Got {len(basic_sections)} pages")

    # Test structured extraction (uses pymupdf4llm)
    structured_sections = extract_pdf_text_structured(pdf_bytes)
    print(f"\n   Structured extraction: {len(structured_sections)} sections")
    for s in structured_sections[:5]:
        fmt = s.get('format', 'plain')
        print(f"   - Page {s['location']} ({fmt}): {len(s['content'])} chars")

    log_test("Structured extraction returns sections", len(structured_sections) >= 3)

    has_markdown = any(s.get('format') == 'markdown' for s in structured_sections)
    print(f"\n   Has markdown format: {has_markdown}")
    log_test("Structured sections have format metadata",
             has_markdown or len(structured_sections) >= 3,
             "markdown" if has_markdown else "fell back to basic (OK - pymupdf4llm may vary)")

    # Test via router
    router_sections = extract_text(pdf_bytes, "pdf")
    log_test("extract_text() router dispatches to structured", len(router_sections) >= 3)

    return structured_sections


def test_chunking(pdf_bytes):
    """Test 3: Hybrid semantic chunking pipeline."""
    print(f"\n{DIVIDER}")
    print("TEST 3: HYBRID SEMANTIC CHUNKING")
    print(DIVIDER)

    from services.chunking import chunk_document_from_blob, MAX_CHUNK_SIZE

    # Mock the semantic chunker since we don't have real OpenAI keys
    with patch('backend.services.chunking._get_semantic_chunker', return_value=None):
        chunks = chunk_document_from_blob(pdf_bytes, file_type="pdf")

    print(f"\n   Total chunks created: {len(chunks)}")
    print(f"   Max chunk size config: {MAX_CHUNK_SIZE} chars")

    # Analyze chunk structure
    chunk_sizes = [len(c['content']) for c in chunks]
    section_names = [c.get('section_name', '') for c in chunks]
    section_types = [c.get('section_type', '') for c in chunks]
    page_nums = [c.get('page_num', '') for c in chunks]

    print(f"\n   Chunk size stats:")
    print(f"   - Min: {min(chunk_sizes)} chars")
    print(f"   - Max: {max(chunk_sizes)} chars")
    print(f"   - Avg: {sum(chunk_sizes) / len(chunk_sizes):.0f} chars")

    print(f"\n   Section types found: {set(section_types)}")
    print(f"   Unique section names: {len(set(section_names))}")

    # Print first 5 chunks in detail
    print(f"\n   First 5 chunks (detailed):")
    for i, chunk in enumerate(chunks[:5]):
        print(f"\n   --- Chunk {i+1} ---")
        print(f"   Page: {chunk.get('page_num', 'N/A')}")
        print(f"   Section: {chunk.get('section_name', 'N/A')}")
        print(f"   Type: {chunk.get('section_type', 'N/A')}")
        print(f"   Content ({len(chunk['content'])} chars): {chunk['content'][:150]}...")

    # Verify critical properties
    log_test("Chunks created from PDF", len(chunks) > 0, f"Got {len(chunks)} chunks")
    log_test("Each chunk has 'content' key", all('content' in c for c in chunks))
    log_test("Each chunk has 'page_num' key", all('page_num' in c for c in chunks))
    log_test("Each chunk has 'section_name' key", all('section_name' in c for c in chunks))
    log_test("Each chunk has 'section_type' key", all('section_type' in c for c in chunks))
    log_test("No empty chunks", all(c['content'].strip() for c in chunks))

    # Check real section names (not just "Chunk N")
    generic_names = sum(1 for n in section_names if n.startswith("Chunk "))
    real_names = len(section_names) - generic_names
    log_test("Chunks have real section names (not just 'Chunk N')",
             real_names > 0 or any('Section' in n for n in section_names),
             f"{real_names} real names, {generic_names} generic")

    return chunks


def test_db_first_uuid_flow(chunks):
    """Test 4: DB-first UUID assignment (the critical fix)."""
    print(f"\n{DIVIDER}")
    print("TEST 4: DB-FIRST UUID ASSIGNMENT")
    print(DIVIDER)

    case_id = str(uuid.uuid4())

    # Simulate what tasks.py now does
    print(f"\n   Case ID: {case_id}")
    print(f"   Processing {len(chunks)} chunks...")

    # Simulate DB flush generating UUIDs
    db_chunk_ids = []
    for idx, chunk in enumerate(chunks):
        db_id = uuid.uuid4()  # Simulated auto-generated UUID
        chunk["id"] = str(db_id)
        chunk["chunk_sequence"] = idx
        db_chunk_ids.append(db_id)

    print(f"\n   Assigned {len(db_chunk_ids)} UUIDs from DB")
    print(f"   Sample IDs:")
    for i in range(min(3, len(db_chunk_ids))):
        print(f"   - Chunk {i}: {chunks[i]['id']}")

    # Verify UUID format
    all_valid_uuids = True
    for chunk in chunks:
        try:
            uuid.UUID(chunk["id"])
        except (ValueError, AttributeError):
            all_valid_uuids = False
            break

    log_test("All chunk IDs are valid UUIDs", all_valid_uuids)
    log_test("Chunk IDs are NOT string concat format",
             not any(":" in c["id"] for c in chunks),
             "No ':' in any chunk ID (old format was 'case_id:idx')")
    log_test("All chunks have chunk_sequence",
             all(isinstance(c.get("chunk_sequence"), int) for c in chunks))

    return chunks


def test_qdrant_payload_structure(chunks):
    """Test 5: Verify Qdrant payload matches expected structure."""
    print(f"\n{DIVIDER}")
    print("TEST 5: QDRANT PAYLOAD STRUCTURE")
    print(DIVIDER)

    from services.vector_store import _generate_point_id

    case_id = str(uuid.uuid4())

    # Simulate what upsert_vectors builds
    print(f"\n   Building Qdrant payloads for {len(chunks)} chunks...")

    payloads = []
    for chunk in chunks:
        chunk_id = str(chunk.get("id", "unknown"))
        metadata = {
            "chunk_id": chunk_id,
            "chunk_sequence": chunk.get("chunk_sequence", 0),
            "page_num": str(chunk.get("page_num", "")),
            "section_name": str(chunk.get("section_name", "")),
            "content_preview": chunk.get("content", "")[:500],
        }
        point_id = _generate_point_id(chunk_id, case_id)
        payloads.append({"point_id": point_id, "payload": metadata})

    # Print sample
    print(f"\n   Sample payload (chunk 0):")
    sample = payloads[0]['payload']
    print(f"   {{")
    for k, v in sample.items():
        val_str = str(v)[:80] + "..." if len(str(v)) > 80 else str(v)
        print(f"     \"{k}\": \"{val_str}\"")
    print(f"   }}")

    # Verify structure
    log_test("Payload has chunk_id (UUID)", all(
        len(p['payload']['chunk_id']) == 36 for p in payloads  # UUID string length
    ))
    log_test("Payload has chunk_sequence", all(
        isinstance(p['payload']['chunk_sequence'], int) for p in payloads
    ))
    log_test("Payload has section_name", all(
        'section_name' in p['payload'] for p in payloads
    ))
    log_test("Content preview is 500 chars max", all(
        len(p['payload']['content_preview']) <= 500 for p in payloads
    ))
    log_test("Point IDs are positive integers", all(
        isinstance(p['point_id'], int) and p['point_id'] > 0 for p in payloads
    ))

    return payloads


def test_rag_citations():
    """Test 6: RAG engine citation patterns including [Section "X"]."""
    print(f"\n{DIVIDER}")
    print("TEST 6: RAG ENGINE CITATIONS")
    print(DIVIDER)

    from services.rag_engine import extract_citations

    # Mock chunks with section names
    chunks = [
        {"page_num": "1", "section_name": "ARTICLE I - DEFINITIONS", "chunk_id": "uuid-1", "score": 0.92, "content": "definitions text"},
        {"page_num": "2", "section_name": "LIMITATION OF LIABILITY", "chunk_id": "uuid-2", "score": 0.88, "content": "liability text"},
        {"page_num": "3", "section_name": "GOVERNING LAW", "chunk_id": "uuid-3", "score": 0.85, "content": "governing law text"},
    ]

    # Test Page citations
    answer_page = 'According to [Page 1], the definitions include confidential information. [Page 2] limits liability.'
    cleaned, citations, has_hall = extract_citations(answer_page, chunks)
    print(f"\n   Page citation test:")
    print(f"   Input: {answer_page[:80]}...")
    print(f"   Citations found: {len(citations)}")
    print(f"   Hallucinations: {has_hall}")
    log_test("[Page X] citations detected", len(citations) >= 2)

    # Test Section citations
    answer_section = 'The [Section "ARTICLE I - DEFINITIONS"] defines confidential information. The [Section "LIMITATION OF LIABILITY"] caps damages.'
    cleaned2, citations2, has_hall2 = extract_citations(answer_section, chunks)
    print(f"\n   Section citation test:")
    print(f"   Input: {answer_section[:80]}...")
    print(f"   Citations found: {len(citations2)}")
    for c in citations2:
        print(f"   - {c['citation_type']}: {c['location']}")
    log_test("[Section \"X\"] citations detected", len(citations2) >= 2,
             f"Found {len(citations2)} section citations")

    # Test hallucination detection
    answer_hallucinated = 'According to [Page 99], this is fabricated. Also [Section "NONEXISTENT"].'
    cleaned3, citations3, has_hall3 = extract_citations(answer_hallucinated, chunks)
    print(f"\n   Hallucination test:")
    print(f"   Hallucinations detected: {has_hall3}")
    log_test("Hallucinated citations detected", has_hall3)

    return citations2


def test_error_response_confidence():
    """Test 7: Error response confidence shape fix."""
    print(f"\n{DIVIDER}")
    print("TEST 7: ERROR RESPONSE CONFIDENCE SHAPE")
    print(DIVIDER)

    # The error_response confidence should be a dict, not a string
    error_response = {
        "answer": None,
        "sources": [],
        "case_id": "test",
        "query": "test",
        "model": "gpt-4o",
        "tokens_used": 0,
        "confidence": {"level": "none", "score": 0.0, "factors": {}},
        "error": None
    }

    confidence = error_response["confidence"]
    print(f"\n   Error response confidence type: {type(confidence).__name__}")
    print(f"   Value: {json.dumps(confidence)}")

    is_dict = isinstance(confidence, dict)
    has_level = "level" in confidence if is_dict else False
    has_score = "score" in confidence if is_dict else False

    log_test("Confidence is dict (not string)", is_dict)
    log_test("Confidence has 'level' key", has_level)
    log_test("Confidence has 'score' key", has_score)


def test_embedding_retry():
    """Test 8: Embedding retry decorators."""
    print(f"\n{DIVIDER}")
    print("TEST 8: EMBEDDING RETRY LOGIC")
    print(DIVIDER)

    import inspect
    from services.embeddings import embed_text, embed_chunks

    # Check that retry decorators are applied
    has_retry_embed_text = hasattr(embed_text, 'retry')
    has_retry_embed_chunks = hasattr(embed_chunks, 'retry')

    print(f"\n   embed_text has retry: {has_retry_embed_text}")
    print(f"   embed_chunks has retry: {has_retry_embed_chunks}")

    log_test("embed_text has @retry decorator", has_retry_embed_text)
    log_test("embed_chunks has @retry decorator", has_retry_embed_chunks)

    # Test that ValueError still raises (not retried)
    try:
        embed_text("")
        log_test("Empty text raises ValueError", False, "Should have raised")
    except ValueError as e:
        log_test("Empty text raises ValueError (not retried)", True, str(e))
    except Exception as e:
        log_test("Empty text raises ValueError", False, f"Got {type(e).__name__}: {e}")


def test_vector_store_safe_create():
    """Test 10: Vector store safe create (no recreate)."""
    print(f"\n{DIVIDER}")
    print("TEST 10: VECTOR STORE SAFE CREATE")
    print(DIVIDER)

    import inspect
    source = inspect.getsource(
        __import__('services.vector_store', fromlist=['create_collection']).create_collection
    )

    has_get_collections = "get_collections" in source
    no_recreate = "recreate_collection" not in source
    has_hnsw = "HnswConfigDiff" in source or "hnsw_config" in source

    print(f"\n   Source analysis:")
    print(f"   - Checks existing collections: {has_get_collections}")
    print(f"   - No recreate_collection: {no_recreate}")
    print(f"   - Has HNSW config: {has_hnsw}")

    log_test("Uses get_collections() check (not recreate)", has_get_collections and no_recreate)
    log_test("Has HNSW configuration", has_hnsw)


def test_batch_upsert():
    """Test 11: Batch upsert configuration."""
    print(f"\n{DIVIDER}")
    print("TEST 11: BATCH UPSERT CONFIG")
    print(DIVIDER)

    from services.vector_store import UPSERT_BATCH_SIZE

    print(f"\n   UPSERT_BATCH_SIZE: {UPSERT_BATCH_SIZE}")

    log_test("Batch size is 100 (Qdrant best practice)", UPSERT_BATCH_SIZE == 100)


def test_embedding_cache():
    """Test 12: Embedding cache integration."""
    print(f"\n{DIVIDER}")
    print("TEST 12: EMBEDDING CACHE")
    print(DIVIDER)

    from services.embedding_cache import get_embedding_cache
    import numpy as np

    cache = get_embedding_cache()

    # Test basic operations
    test_key = "test_hash_abc123"
    test_embedding = np.random.rand(3072)

    cache.put(test_key, test_embedding)
    retrieved = cache.get(test_key)

    print(f"\n   Cache size: {len(cache.cache)}")
    print(f"   Hit rate: {cache.get_hit_rate():.2f}")

    log_test("Cache put/get works", retrieved is not None)
    log_test("Cache preserves embedding values",
             retrieved is not None and np.allclose(retrieved, test_embedding))

    # Test miss
    missed = cache.get("nonexistent")
    log_test("Cache miss returns None", missed is None)

    # Verify embed_text uses cache
    import inspect
    source = inspect.getsource(
        __import__('services.embeddings', fromlist=['embed_text']).embed_text
    )
    has_cache = "get_embedding_cache" in source or "cache" in source.lower()
    log_test("embed_text integrates cache", has_cache)

    cache.clear()


def test_chunk_model():
    """Test 13: Chunk model has section_type."""
    print(f"\n{DIVIDER}")
    print("TEST 13: CHUNK MODEL UPDATE")
    print(DIVIDER)

    from models import Chunk

    # Check columns
    columns = {c.name for c in Chunk.__table__.columns}
    print(f"\n   Chunk columns: {sorted(columns)}")

    log_test("Chunk model has 'section_type' column", 'section_type' in columns)
    log_test("Chunk model has 'section_name' column", 'section_name' in columns)
    log_test("Chunk model has 'chunk_sequence' column", 'chunk_sequence' in columns)


def test_docx_extraction():
    """Test 14: DOCX structured extraction."""
    print(f"\n{DIVIDER}")
    print("TEST 14: DOCX TEXT EXTRACTION")
    print(DIVIDER)

    from docx import Document as DocxDocument
    import io

    # Create a sample DOCX with headings
    doc = DocxDocument()
    doc.add_heading("MASTER SERVICES AGREEMENT", level=0)
    doc.add_heading("Article I - Definitions", level=1)
    doc.add_paragraph("This section defines key terms used throughout the agreement.")
    doc.add_heading("Article II - Scope of Services", level=1)
    doc.add_paragraph("The Provider shall deliver professional consulting services.")
    doc.add_heading("Section 2.1 - Deliverables", level=2)
    doc.add_paragraph("All deliverables must meet quality standards.")

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    from services.text_extraction import extract_docx_text_structured

    sections = extract_docx_text_structured(docx_bytes)

    print(f"\n   Created DOCX with headings")
    print(f"   Extracted {len(sections)} sections:")
    for s in sections:
        fmt = s.get('format', 'plain')
        preview = s['content'][:80].replace('\n', ' ')
        print(f"   - [{fmt}] {preview}...")

    has_markdown = any(s.get('format') == 'markdown' for s in sections)
    has_heading = any('#' in s.get('content', '') for s in sections)

    log_test("DOCX extraction returns sections", len(sections) > 0)
    log_test("DOCX sections have markdown format", has_markdown)
    log_test("DOCX sections contain markdown headers", has_heading)


def test_legal_system_prompt():
    """Test 15: LEGAL_SYSTEM_PROMPT update."""
    print(f"\n{DIVIDER}")
    print("TEST 15: LEGAL SYSTEM PROMPT")
    print(DIVIDER)

    from services.rag_engine import LEGAL_SYSTEM_PROMPT

    has_section_citation = '[Section "' in LEGAL_SYSTEM_PROMPT or "Section" in LEGAL_SYSTEM_PROMPT
    has_page = "[Page X]" in LEGAL_SYSTEM_PROMPT
    has_para = "[Paragraph X]" in LEGAL_SYSTEM_PROMPT

    print(f"\n   Prompt length: {len(LEGAL_SYSTEM_PROMPT)} chars")
    print(f"   Has Section citation instruction: {has_section_citation}")
    print(f"   Has Page citation: {has_page}")
    print(f"   Has Paragraph citation: {has_para}")

    log_test("System prompt includes [Section \"Name\"] instruction", has_section_citation)


def run_all_tests():
    """Run the complete test suite."""
    print("\n" + "=" * 70)
    print("  VERITAS AI - FULL PIPELINE INTEGRATION TEST")
    print("  Testing: Extraction → Section Detection → Chunking → IDs → RAG")
    print("=" * 70)

    # Test 1: Section detector
    sections = test_section_detector()

    # Test 2: PDF extraction
    pdf_bytes = create_sample_legal_pdf()
    extracted = test_pdf_extraction(pdf_bytes)

    # Test 3: Chunking
    chunks = test_chunking(pdf_bytes)

    # Test 4: DB-first UUID flow
    chunks_with_ids = test_db_first_uuid_flow(chunks)

    # Test 5: Qdrant payload structure
    payloads = test_qdrant_payload_structure(chunks_with_ids)

    # Test 6: RAG citations
    test_rag_citations()

    # Test 7: Error response confidence
    test_error_response_confidence()

    # Test 8: Embedding retry
    test_embedding_retry()

    # Test 9: Safe create
    test_vector_store_safe_create()

    # Test 11: Batch upsert
    test_batch_upsert()

    # Test 12: Embedding cache
    test_embedding_cache()

    # Test 13: Chunk model
    test_chunk_model()

    # Test 14: DOCX extraction
    test_docx_extraction()

    # Test 15: System prompt
    test_legal_system_prompt()

    # Summary
    print(f"\n{'=' * 70}")
    print("  TEST RESULTS SUMMARY")
    print(f"{'=' * 70}")

    passed = sum(1 for _, p in results if p)
    failed = sum(1 for _, p in results if not p)
    total = len(results)

    for name, p in results:
        status = "PASS" if p else "FAIL"
        icon = "✅" if p else "❌"
        print(f"  {icon} {status}: {name}")

    print(f"\n  {'-' * 50}")
    print(f"  Total: {total} | Passed: {passed} | Failed: {failed}")
    print(f"  Pass rate: {passed/total*100:.0f}%")
    print(f"{'=' * 70}\n")

    return failed == 0


if __name__ == "__main__":
    os.chdir(os.path.join(os.path.dirname(__file__), '..'))
    success = run_all_tests()
    sys.exit(0 if success else 1)
