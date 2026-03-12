"""Multi-format text extraction service for PDFs, DOCX, and TXT files

Supports:
- pymupdf4llm for structured PDF extraction (markdown with headers)
- python-docx with heading detection for DOCX files
- Basic extraction for TXT files

TODO: Evaluate replacing pymupdf4llm with Docling (pip install docling)
https://github.com/docling-project/docling — IBM Research open-source.
Docling uses DocLayNet (ML layout analysis) to detect headings, paragraphs,
tables, footnotes as a structured hierarchy. This would provide real heading
detection instead of relying on markdown header heuristics. Benchmark shows
97.9% accuracy on complex documents vs pymupdf4llm's basic text extraction.
Key benefit: eliminates the regex-based _extract_section_label() in chunking.py
since Docling natively provides heading/paragraph structure.
"""
import logging
import tempfile
import os
from typing import List, Dict
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


def extract_pdf_text_structured(file_bytes: bytes) -> List[Dict[str, str]]:
    """
    Extract structured text from PDF using pymupdf4llm.

    Produces markdown-formatted output with headers detected from font sizes.
    Falls back to basic extraction on failure.

    Args:
        file_bytes: Raw PDF bytes

    Returns:
        List of dicts with keys: content, location, location_type, format
    """
    if not file_bytes:
        raise ValueError("PDF content is empty")

    try:
        import pymupdf4llm

        logger.info(f"Extracting structured text from PDF ({len(file_bytes)} bytes)")
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        # Use pymupdf4llm for markdown extraction with page chunks
        page_chunks = pymupdf4llm.to_markdown(
            doc,
            page_chunks=True,
        )

        if not page_chunks:
            logger.warning("pymupdf4llm returned empty result, falling back")
            doc.close()
            return extract_pdf_text(file_bytes)

        sections = []
        for i, chunk in enumerate(page_chunks):
            text = chunk.get("text", "") if isinstance(chunk, dict) else str(chunk)
            if text.strip():
                sections.append({
                    "content": text,
                    "location": str(i + 1),
                    "location_type": "page",
                    "format": "markdown"
                })

        doc.close()
        logger.info(f"Extracted {len(sections)} structured pages from PDF")
        return sections

    except ImportError:
        logger.warning("pymupdf4llm not available, falling back to basic extraction")
        return extract_pdf_text(file_bytes)
    except Exception as e:
        logger.warning(f"Structured PDF extraction failed ({str(e)}), falling back to basic")
        return extract_pdf_text(file_bytes)


def extract_pdf_text(file_bytes: bytes) -> List[Dict[str, str]]:
    """
    Extract text from PDF file using PyMuPDF (fitz).

    Args:
        file_bytes: Raw PDF bytes

    Returns:
        List of dicts with keys: content, location (page number), location_type

    Raises:
        ValueError: If PDF is invalid or empty
        Exception: If PDF parsing fails
    """
    if not file_bytes:
        raise ValueError("PDF content is empty")

    try:
        logger.info(f"Extracting text from PDF ({len(file_bytes)} bytes)")
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")

        if pdf_document.page_count == 0:
            raise ValueError("PDF contains no pages or failed to load")

        logger.info(f"Loaded {pdf_document.page_count} pages from PDF")

        sections = []
        for page_idx in range(pdf_document.page_count):
            page = pdf_document[page_idx]
            page_text = page.get_text()

            if page_text.strip():
                sections.append({
                    "content": page_text,
                    "location": str(page_idx + 1),  # 1-indexed page number
                    "location_type": "page"
                })

        pdf_document.close()
        logger.info(f"Extracted {len(sections)} pages from PDF")
        return sections

    except Exception as e:
        logger.error(f"Error extracting PDF text: {str(e)}")
        raise


def extract_docx_text_structured(file_bytes: bytes) -> List[Dict[str, str]]:
    """
    Extract heading-aware text from DOCX using python-docx.

    Detects Word heading styles and converts to markdown-formatted output.
    Falls back to basic extraction on failure.

    Args:
        file_bytes: Raw DOCX bytes

    Returns:
        List of dicts with keys: content, location, location_type, format
    """
    if not file_bytes:
        raise ValueError("DOCX content is empty")

    temp_file = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            temp_file = tmp.name

        logger.info(f"Extracting structured text from DOCX ({len(file_bytes)} bytes)")
        doc = DocxDocument(temp_file)

        if not doc.paragraphs:
            raise ValueError("DOCX contains no paragraphs or failed to load")

        # Build markdown output grouping by headings
        sections = []
        current_section_lines = []
        para_counter = 0

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            para_counter += 1
            style_name = para.style.name if para.style else ""

            # Detect heading levels
            if style_name.startswith("Heading"):
                # Flush current section
                if current_section_lines:
                    sections.append({
                        "content": "\n".join(current_section_lines),
                        "location": f"para {para_counter}",
                        "location_type": "paragraph",
                        "format": "markdown"
                    })
                    current_section_lines = []

                # Determine heading level
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                md_prefix = "#" * min(level, 6)
                current_section_lines.append(f"{md_prefix} {para_text}")
            elif style_name == "Title":
                if current_section_lines:
                    sections.append({
                        "content": "\n".join(current_section_lines),
                        "location": f"para {para_counter}",
                        "location_type": "paragraph",
                        "format": "markdown"
                    })
                    current_section_lines = []
                current_section_lines.append(f"# {para_text}")
            else:
                current_section_lines.append(para_text)

        # Flush remaining
        if current_section_lines:
            sections.append({
                "content": "\n".join(current_section_lines),
                "location": f"para {para_counter}",
                "location_type": "paragraph",
                "format": "markdown"
            })

        if not sections:
            logger.warning("No structured sections found in DOCX, falling back")
            return extract_docx_text(file_bytes)

        logger.info(f"Extracted {len(sections)} structured sections from DOCX")
        return sections

    except Exception as e:
        logger.warning(f"Structured DOCX extraction failed ({str(e)}), falling back to basic")
        return extract_docx_text(file_bytes)
    finally:
        if temp_file and os.path.exists(temp_file):
            try:
                os.unlink(temp_file)
            except OSError:
                pass


def extract_docx_text(file_bytes: bytes) -> List[Dict[str, str]]:
    """
    Extract text from DOCX (Word) file using python-docx.

    Args:
        file_bytes: Raw DOCX bytes

    Returns:
        List of dicts with keys: content, location (paragraph number), location_type

    Raises:
        ValueError: If DOCX is invalid or empty
        Exception: If DOCX parsing fails
    """
    if not file_bytes:
        raise ValueError("DOCX content is empty")

    temp_file = None
    try:
        # Write to temporary file for python-docx processing
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            temp_file = tmp.name

        logger.info(f"Extracting text from DOCX ({len(file_bytes)} bytes)")
        doc = DocxDocument(temp_file)

        if not doc.paragraphs:
            raise ValueError("DOCX contains no paragraphs or failed to load")

        logger.info(f"Loaded {len(doc.paragraphs)} paragraphs from DOCX")

        sections = []
        para_counter = 0

        for para in doc.paragraphs:
            para_text = para.text.strip()

            # Skip empty paragraphs but count them for location tracking
            if para_text:
                para_counter += 1
                sections.append({
                    "content": para_text,
                    "location": f"para {para_counter}",  # Paragraph number
                    "location_type": "paragraph"
                })

        logger.info(f"Extracted {len(sections)} non-empty paragraphs from DOCX")
        return sections

    except Exception as e:
        logger.error(f"Error extracting DOCX text: {str(e)}")
        raise
    finally:
        # Clean up temporary file
        if temp_file and os.path.exists(temp_file):
            try:
                os.unlink(temp_file)
                logger.debug(f"Cleaned up temporary file: {temp_file}")
            except OSError as e:
                logger.warning(f"Failed to delete temp file {temp_file}: {str(e)}")


def extract_txt_text(file_bytes: bytes, lines_per_section: int = 50) -> List[Dict[str, str]]:
    """
    Extract text from TXT file.

    Args:
        file_bytes: Raw TXT bytes
        lines_per_section: Number of lines per section (default 50)

    Returns:
        List of dicts with keys: content, location (line range), location_type

    Raises:
        ValueError: If TXT is invalid or empty
        UnicodeDecodeError: If file is not valid UTF-8 text
    """
    if not file_bytes:
        raise ValueError("TXT content is empty")

    try:
        # Decode bytes to string
        text_content = file_bytes.decode('utf-8')
        logger.info(f"Extracting text from TXT ({len(file_bytes)} bytes)")

        lines = text_content.split('\n')
        logger.info(f"Loaded {len(lines)} lines from TXT")

        sections = []
        for i in range(0, len(lines), lines_per_section):
            section_lines = lines[i:i + lines_per_section]
            section_text = '\n'.join(section_lines).strip()

            if section_text:
                start_line = i + 1  # 1-indexed
                end_line = min(i + lines_per_section, len(lines))
                location = f"line {start_line}-{end_line}"

                sections.append({
                    "content": section_text,
                    "location": location,
                    "location_type": "line_range"
                })

        logger.info(f"Extracted {len(sections)} sections from TXT")
        return sections

    except UnicodeDecodeError as e:
        logger.error(f"TXT file is not valid UTF-8: {str(e)}")
        raise ValueError(f"TXT file must be valid UTF-8 encoded text: {str(e)}") from e
    except Exception as e:
        logger.error(f"Error extracting TXT text: {str(e)}")
        raise


def extract_text(file_bytes: bytes, file_type: str) -> List[Dict[str, str]]:
    """
    Route text extraction to appropriate handler based on file type.

    Uses structured extraction (pymupdf4llm / heading-aware DOCX) by default
    with automatic fallback to basic extraction.

    Args:
        file_bytes: Raw file bytes
        file_type: File type string ('pdf', 'docx', or 'txt')

    Returns:
        List of dicts with content and location information

    Raises:
        ValueError: If file_type is unsupported or content is invalid
        Exception: If extraction fails
    """
    if not file_bytes:
        raise ValueError("File content is empty")

    if file_type == "pdf":
        return extract_pdf_text_structured(file_bytes)
    elif file_type == "docx":
        return extract_docx_text_structured(file_bytes)
    elif file_type == "txt":
        return extract_txt_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
