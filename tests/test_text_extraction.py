"""Tests for multi-format text extraction service with Docling"""
import pytest
import os
from io import BytesIO

# Set required env vars before imports
os.environ.setdefault('DATABASE_URL', 'sqlite:////tmp/test_lexintel.db')
os.environ.setdefault('OPENAI_API_KEY', 'sk-test')
os.environ.setdefault('AZURE_STORAGE_CONNECTION_STRING', 'UseDevelopmentStorage=true')
os.environ.setdefault('SECRET_KEY', 'test-secret-key-for-testing-long-enough')

from backend.services.text_extraction import (
    extract_text,
    extract_pdf_text,
    extract_txt_text,
    extract_docx_text
)


def create_minimal_pdf() -> bytes:
    """Create minimal valid PDF for testing"""
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT
/F1 12 Tf
100 700 Td
(Hello from PDF) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000204 00000 n
trailer
<< /Size 5 /Root 1 0 R >>
startxref
297
%%EOF"""
    return pdf_content


def create_minimal_docx() -> bytes:
    """Create minimal valid DOCX for testing"""
    from io import BytesIO
    from docx import Document

    # Create a proper DOCX using python-docx
    doc = Document()
    doc.add_paragraph("Hello from DOCX")
    doc.add_paragraph("Second paragraph")

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def create_minimal_txt() -> bytes:
    """Create minimal TXT file for testing"""
    return b"""This is a test text file.
It has multiple lines.
Line 3 goes here.
And line 4 as well.
Fifth line of content.
Sixth line here."""


class TestTextExtractionPDF:
    """Test PDF text extraction"""

    def test_extract_pdf_text_valid(self):
        """Test extracting text from valid PDF"""
        pdf_bytes = create_minimal_pdf()
        result = extract_pdf_text(pdf_bytes)

        assert isinstance(result, list)
        assert len(result) > 0
        assert all("content" in item for item in result)
        assert all("location" in item for item in result)
        assert all("location_type" in item for item in result)

    def test_extract_pdf_text_empty(self):
        """Test that empty PDF raises ValueError"""
        with pytest.raises(ValueError, match="empty"):
            extract_pdf_text(b"")

    def test_extract_pdf_text_invalid(self):
        """Test that invalid PDF raises exception"""
        invalid_pdf = b"Not a real PDF"
        with pytest.raises(Exception):
            extract_pdf_text(invalid_pdf)

    def test_pdf_location_format(self):
        """Test that PDF locations are formatted correctly"""
        pdf_bytes = create_minimal_pdf()
        result = extract_pdf_text(pdf_bytes)

        for item in result:
            # Should be page numbers (numeric strings)
            assert item["location"].isdigit()
            assert item["location_type"] == "page"


class TestTextExtractionDOCX:
    """Test DOCX text extraction"""

    def test_extract_docx_text_valid(self):
        """Test extracting text from valid DOCX"""
        docx_bytes = create_minimal_docx()
        result = extract_docx_text(docx_bytes)

        assert isinstance(result, list)
        assert len(result) > 0
        assert all("content" in item for item in result)
        assert all("location" in item for item in result)
        assert all("location_type" in item for item in result)

    def test_extract_docx_text_empty(self):
        """Test that empty DOCX raises ValueError"""
        with pytest.raises(ValueError, match="empty"):
            extract_docx_text(b"")

    def test_extract_docx_text_invalid(self):
        """Test that invalid DOCX raises exception"""
        invalid_docx = b"Not a real DOCX"
        with pytest.raises(Exception):
            extract_docx_text(invalid_docx)

    def test_docx_location_format(self):
        """Test that DOCX locations are formatted correctly"""
        docx_bytes = create_minimal_docx()
        result = extract_docx_text(docx_bytes)

        for item in result:
            # Should be paragraph references (para X format)
            assert item["location"].startswith("para")
            assert item["location_type"] == "paragraph"


class TestTextExtractionTXT:
    """Test TXT text extraction"""

    def test_extract_txt_text_valid(self):
        """Test extracting text from valid TXT"""
        txt_bytes = create_minimal_txt()
        result = extract_txt_text(txt_bytes)

        assert isinstance(result, list)
        assert len(result) > 0
        assert all("content" in item for item in result)
        assert all("location" in item for item in result)
        assert all("location_type" in item for item in result)

    def test_extract_txt_text_empty(self):
        """Test that empty TXT raises ValueError"""
        with pytest.raises(ValueError, match="empty"):
            extract_txt_text(b"")

    def test_extract_txt_text_invalid_encoding(self):
        """Test that non-UTF8 TXT raises ValueError"""
        invalid_txt = b"\x80\x81\x82\x83"  # Invalid UTF-8
        with pytest.raises(ValueError, match="UTF-8"):
            extract_txt_text(invalid_txt)

    def test_txt_location_format(self):
        """Test that TXT locations are formatted correctly"""
        txt_bytes = create_minimal_txt()
        result = extract_txt_text(txt_bytes)

        for item in result:
            # Should be line range references (line X-Y format)
            assert item["location"].startswith("line")
            assert "-" in item["location"]
            assert item["location_type"] == "line_range"

    def test_txt_line_grouping(self):
        """Test that TXT lines are grouped correctly"""
        txt_bytes = create_minimal_txt()
        result = extract_txt_text(txt_bytes, lines_per_section=2)

        # With 6 lines and 2 lines per section, expect 3 sections
        assert len(result) == 3


class TestTextExtractionRouter:
    """Test extract_text router function"""

    def test_extract_text_pdf(self):
        """Test router with PDF type"""
        pdf_bytes = create_minimal_pdf()
        result = extract_text(pdf_bytes, "pdf")

        assert isinstance(result, list)
        assert all(item["location_type"] == "page" for item in result)

    def test_extract_text_docx(self):
        """Test router with DOCX type"""
        docx_bytes = create_minimal_docx()
        result = extract_text(docx_bytes, "docx")

        assert isinstance(result, list)
        assert all(item["location_type"] == "paragraph" for item in result)

    def test_extract_text_txt(self):
        """Test router with TXT type"""
        txt_bytes = create_minimal_txt()
        result = extract_text(txt_bytes, "txt")

        assert isinstance(result, list)
        assert all(item["location_type"] == "line_range" for item in result)

    def test_extract_text_empty(self):
        """Test that empty content raises ValueError"""
        with pytest.raises(ValueError, match="empty"):
            extract_text(b"", "pdf")

    def test_extract_text_unsupported_type(self):
        """Test that unsupported file type raises ValueError"""
        with pytest.raises(ValueError, match="Unsupported"):
            extract_text(create_minimal_pdf(), "docm")


class TestExtractionConsistency:
    """Test that all extractors return consistent format"""

    def test_all_formats_have_required_fields(self):
        """Test that all formats return required fields"""
        pdf_result = extract_pdf_text(create_minimal_pdf())
        docx_result = extract_docx_text(create_minimal_docx())
        txt_result = extract_txt_text(create_minimal_txt())

        required_fields = ["content", "location", "location_type"]

        for result in [pdf_result, docx_result, txt_result]:
            for item in result:
                assert all(field in item for field in required_fields)

    def test_all_formats_have_content(self):
        """Test that all formats extract content"""
        pdf_result = extract_pdf_text(create_minimal_pdf())
        docx_result = extract_docx_text(create_minimal_docx())
        txt_result = extract_txt_text(create_minimal_txt())

        for result in [pdf_result, docx_result, txt_result]:
            assert len(result) > 0
            assert all(len(item["content"]) > 0 for item in result)

    def test_location_types_distinct(self):
        """Test that each format has distinct location type"""
        pdf_result = extract_pdf_text(create_minimal_pdf())
        docx_result = extract_docx_text(create_minimal_docx())
        txt_result = extract_txt_text(create_minimal_txt())

        pdf_location_types = set(item["location_type"] for item in pdf_result)
        docx_location_types = set(item["location_type"] for item in docx_result)
        txt_location_types = set(item["location_type"] for item in txt_result)

        # Should have exactly one location type each
        assert pdf_location_types == {"page"}
        assert docx_location_types == {"paragraph"}
        assert txt_location_types == {"line_range"}
