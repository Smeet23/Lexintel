"""
Integration tests for the complete document extraction pipeline.
Tests the flow: extract_file -> chunk_text -> create_text_chunks
"""
import pytest
from pathlib import Path
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter


@pytest.mark.asyncio
async def test_extraction_pipeline_txt(tmp_path, db_session):
    """Test complete extraction pipeline with TXT file"""
    from app.services.extraction import extract_file, create_text_chunks
    from shared.models import Document, DocumentChunk

    # Create a test document in database
    test_doc_id = "doc_test_001"
    doc = Document(
        id=test_doc_id,
        case_id="case_test_001",
        title="Test Document",
        filename="test.txt",
        type="OTHER",
    )
    db_session.add(doc)
    await db_session.commit()

    # Create a test TXT file
    txt_path = tmp_path / "test.txt"
    test_content = "This is a test document. " * 500  # Make it long enough to chunk
    txt_path.write_text(test_content)

    # Extract text
    extracted_text = extract_file(str(txt_path))
    assert extracted_text == test_content

    # Create chunks
    chunk_ids = await create_text_chunks(
        document_id=test_doc_id,
        text=extracted_text,
        session=db_session,
        chunk_size=200,
        overlap=50,
    )

    # Verify chunks were created
    assert len(chunk_ids) > 1
    assert all(isinstance(cid, str) for cid in chunk_ids)

    # Verify chunks are in database
    from sqlalchemy import select

    stmt = select(DocumentChunk).where(DocumentChunk.document_id == test_doc_id)
    result = await db_session.execute(stmt)
    chunks = result.scalars().all()

    assert len(chunks) == len(chunk_ids)
    for i, chunk in enumerate(chunks):
        assert chunk.chunk_index == i
        assert chunk.document_id == test_doc_id
        assert len(chunk.chunk_text) > 0


@pytest.mark.asyncio
async def test_extraction_pipeline_pdf(tmp_path, db_session):
    """Test complete extraction pipeline with PDF file"""
    from app.services.extraction import extract_file, create_text_chunks
    from shared.models import Document, DocumentChunk

    # Create a test document in database
    test_doc_id = "doc_test_002"
    doc = Document(
        id=test_doc_id,
        case_id="case_test_002",
        title="Test PDF",
        filename="test.pdf",
        type="OTHER",
    )
    db_session.add(doc)
    await db_session.commit()

    # Create a test PDF file
    pdf_path = tmp_path / "test.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    for i in range(3):
        c.drawString(100, 750 - (i * 20), f"This is line {i + 1} of test content.")
    c.save()

    # Extract text
    extracted_text = extract_file(str(pdf_path))
    assert isinstance(extracted_text, str)
    assert len(extracted_text) > 0

    # Create chunks
    chunk_ids = await create_text_chunks(
        document_id=test_doc_id,
        text=extracted_text,
        session=db_session,
        chunk_size=100,
        overlap=20,
    )

    # Verify chunks were created
    assert len(chunk_ids) >= 1

    # Verify chunks are in database
    from sqlalchemy import select

    stmt = select(DocumentChunk).where(DocumentChunk.document_id == test_doc_id)
    result = await db_session.execute(stmt)
    chunks = result.scalars().all()

    assert len(chunks) == len(chunk_ids)


@pytest.mark.asyncio
async def test_extraction_pipeline_docx(tmp_path, db_session):
    """Test complete extraction pipeline with DOCX file"""
    from app.services.extraction import extract_file, create_text_chunks
    from shared.models import Document, DocumentChunk

    # Create a test document in database
    test_doc_id = "doc_test_003"
    doc = Document(
        id=test_doc_id,
        case_id="case_test_003",
        title="Test DOCX",
        filename="test.docx",
        type="OTHER",
    )
    db_session.add(doc)
    await db_session.commit()

    # Create a test DOCX file
    docx_path = tmp_path / "test.docx"
    doc_word = DocxDocument()
    for i in range(5):
        doc_word.add_paragraph(f"This is paragraph {i + 1} with test content.")
    doc_word.save(str(docx_path))

    # Extract text
    extracted_text = extract_file(str(docx_path))
    assert isinstance(extracted_text, str)
    assert "paragraph" in extracted_text.lower()

    # Create chunks
    chunk_ids = await create_text_chunks(
        document_id=test_doc_id,
        text=extracted_text,
        session=db_session,
        chunk_size=150,
        overlap=30,
    )

    # Verify chunks were created
    assert len(chunk_ids) >= 1

    # Verify chunks are in database
    from sqlalchemy import select

    stmt = select(DocumentChunk).where(DocumentChunk.document_id == test_doc_id)
    result = await db_session.execute(stmt)
    chunks = result.scalars().all()

    assert len(chunks) == len(chunk_ids)


@pytest.mark.asyncio
async def test_extraction_pipeline_overlap(tmp_path, db_session):
    """Test that chunks have proper overlap"""
    from app.services.extraction import extract_file, create_text_chunks
    from shared.models import Document, DocumentChunk

    # Create a test document in database
    test_doc_id = "doc_test_overlap"
    doc = Document(
        id=test_doc_id,
        case_id="case_test_overlap",
        title="Overlap Test",
        filename="test.txt",
        type="OTHER",
    )
    db_session.add(doc)
    await db_session.commit()

    # Create a test file with known content
    txt_path = tmp_path / "overlap_test.txt"
    test_content = "ABCDEFGHIJ" * 50  # 500 characters
    txt_path.write_text(test_content)

    # Extract and create chunks with specific sizes
    extracted_text = extract_file(str(txt_path))

    chunk_ids = await create_text_chunks(
        document_id=test_doc_id,
        text=extracted_text,
        session=db_session,
        chunk_size=100,
        overlap=20,
    )

    # Verify chunks
    from sqlalchemy import select

    stmt = select(DocumentChunk).where(DocumentChunk.document_id == test_doc_id)
    result = await db_session.execute(stmt)
    chunks = result.scalars().all()

    # Verify overlap between consecutive chunks
    for i in range(len(chunks) - 1):
        current_chunk = chunks[i].chunk_text
        next_chunk = chunks[i + 1].chunk_text

        # The end of current chunk should overlap with start of next chunk
        overlap_content = current_chunk[-20:]  # Last 20 chars of current
        assert next_chunk.startswith(overlap_content), (
            f"Overlap mismatch between chunk {i} and {i+1}"
        )
